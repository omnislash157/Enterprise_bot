"""
Document Loader - Multi-format loading and context stuffing for enterprise mode.

Loads .docx, .json, .csv, .xlsx, .md, .txt files from hierarchical folder structure,
caches content, and provides division-aware context building.

Folder structure:
    manuals/
    ├── Driscoll/
    │   ├── Warehouse/        # Division folder
    │   │   ├── Dispatching Manual.docx
    │   │   ├── Driver Check-in Manual.docx
    │   │   └── ...
    │   ├── Sales/
    │   │   ├── telnet_sop_chunks.json   # JSON chunk files work too
    │   │   └── ...
    │   ├── HR/
    │   ├── Purchasing/
    │   └── Shared/           # Accessible by all divisions

Usage:
    from doc_loader import DocLoader, DivisionContextBuilder

    loader = DocLoader(Path("./manuals"))
    stats = loader.get_stats()

    builder = DivisionContextBuilder(loader)
    context = builder.get_context_for_division("warehouse", max_tokens=200000)

Version: 2.0.0 (multi-format)
"""

import json
import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# Supported file extensions
SUPPORTED_EXTENSIONS = {".docx", ".json", ".csv", ".xlsx", ".md", ".txt"}

logger = logging.getLogger(__name__)


# =============================================================================
# DATA STRUCTURES
# =============================================================================

@dataclass
class LoadedDoc:
    """A loaded document with metadata."""
    path: Path
    name: str
    division: str
    content: str
    char_count: int
    approx_tokens: int
    paragraphs: int

    def to_dict(self) -> Dict[str, Any]:
        return {
            "name": self.name,
            "division": self.division,
            "char_count": self.char_count,
            "approx_tokens": self.approx_tokens,
            "paragraphs": self.paragraphs,
        }


@dataclass
class DocStats:
    """Statistics about loaded documents."""
    total_docs: int
    total_chars: int
    total_tokens: int
    by_division: Dict[str, Dict[str, int]]
    doc_list: List[str]


# =============================================================================
# DOCX LOADER
# =============================================================================

class DocLoader:
    """
    Loads and caches document files from a directory tree.

    Supports: .docx, .json, .csv, .xlsx, .md, .txt
    Extracts text content and organizes by division (folder structure).
    """

    # Approximate tokens per character (conservative estimate)
    CHARS_PER_TOKEN = 4

    def __init__(self, docs_dir: Path):
        """
        Initialize document loader.

        Args:
            docs_dir: Root directory containing documents
        """
        self.docs_dir = Path(docs_dir)
        self._cache: Dict[str, LoadedDoc] = {}
        self._loaded = False

    def _extract_text(self, docx_path: Path) -> str:
        """Extract all text from a .docx file."""
        if not DOCX_AVAILABLE:
            logger.warning(f"python-docx not available, skipping {docx_path}")
            return ""
        try:
            doc = Document(docx_path)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            return "\n\n".join(paragraphs)

        except Exception as e:
            logger.error(f"Error extracting text from {docx_path}: {e}")
            return ""

    def _load_json(self, file_path: Path) -> str:
        """Load JSON file - handles chunk arrays or plain objects."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)

            # Handle chunk array format (from upload_manuals.py)
            if isinstance(data, list):
                chunks = []
                for item in data:
                    if isinstance(item, dict):
                        # Chunk format: {"chunk_index": N, "content": "..."}
                        content = item.get("content", "")
                        if content:
                            chunks.append(content)
                    elif isinstance(item, str):
                        chunks.append(item)
                return "\n\n".join(chunks)

            # Plain object - convert to readable text
            elif isinstance(data, dict):
                lines = []
                for key, value in data.items():
                    lines.append(f"{key}: {value}")
                return "\n".join(lines)

            return str(data)

        except Exception as e:
            logger.error(f"Error loading JSON {file_path}: {e}")
            return ""

    def _load_csv(self, file_path: Path) -> str:
        """Load CSV file as formatted text."""
        if not PANDAS_AVAILABLE:
            # Fallback: read as plain text
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            except Exception as e:
                logger.error(f"Error loading CSV {file_path}: {e}")
                return ""

        try:
            df = pd.read_csv(file_path)
            return df.to_string(index=False)
        except Exception as e:
            logger.error(f"Error loading CSV {file_path}: {e}")
            return ""

    def _load_excel(self, file_path: Path) -> str:
        """Load Excel file as formatted text."""
        if not PANDAS_AVAILABLE:
            logger.warning(f"pandas not available, skipping Excel {file_path}")
            return ""

        try:
            # Read all sheets
            xlsx = pd.ExcelFile(file_path)
            sheets_text = []

            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                sheet_text = f"=== Sheet: {sheet_name} ===\n{df.to_string(index=False)}"
                sheets_text.append(sheet_text)

            return "\n\n".join(sheets_text)

        except Exception as e:
            logger.error(f"Error loading Excel {file_path}: {e}")
            return ""

    def _load_text(self, file_path: Path) -> str:
        """Load plain text file (.txt, .md)."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            return ""

    def _load_file(self, file_path: Path) -> str:
        """Load file based on extension - dispatcher method."""
        ext = file_path.suffix.lower()

        if ext == ".docx":
            return self._extract_text(file_path)
        elif ext == ".json":
            return self._load_json(file_path)
        elif ext == ".csv":
            return self._load_csv(file_path)
        elif ext == ".xlsx":
            return self._load_excel(file_path)
        elif ext in (".txt", ".md"):
            return self._load_text(file_path)
        else:
            logger.warning(f"Unsupported file type: {ext} for {file_path}")
            return ""

    def _detect_division(self, docx_path: Path) -> str:
        """
        Detect division from folder path.

        Folder structure relative to docs_dir:
            Division/foo.docx -> "division"

        Examples (docs_dir = manuals/Driscoll):
            Warehouse/foo.docx -> "warehouse"
            HR/bar.docx -> "hr"
            Shared/baz.docx -> "shared"
        """
        # Get path relative to docs_dir
        try:
            rel_path = docx_path.relative_to(self.docs_dir)
            parts = rel_path.parts
        except ValueError:
            return "general"

        # parts[0] = division folder, parts[-1] = filename
        if len(parts) >= 2:
            # File is in a subfolder - that's the division
            division = parts[0].lower()
            return division
        elif len(parts) == 1:
            # File directly in docs_dir (no division folder)
            return "general"

        return "general"

    def _load_all(self):
        """Load all supported files from docs directory."""
        if self._loaded:
            return

        if not self.docs_dir.exists():
            logger.warning(f"Docs directory not found: {self.docs_dir}")
            self._loaded = True
            return

        # Find all supported files recursively
        all_files = []
        for ext in SUPPORTED_EXTENSIONS:
            pattern = f"*{ext}"
            all_files.extend(self.docs_dir.rglob(pattern))

        # Filter out temp files (start with ~)
        all_files = [f for f in all_files if not f.name.startswith("~")]

        logger.info(f"Found {len(all_files)} files in {self.docs_dir}")

        for file_path in all_files:
            try:
                content = self._load_file(file_path)
                if not content:
                    continue

                division = self._detect_division(file_path)
                char_count = len(content)
                approx_tokens = char_count // self.CHARS_PER_TOKEN
                para_count = content.count("\n\n") + 1

                doc = LoadedDoc(
                    path=file_path,
                    name=file_path.stem,  # Filename without extension
                    division=division,
                    content=content,
                    char_count=char_count,
                    approx_tokens=approx_tokens,
                    paragraphs=para_count,
                )

                # Use relative path as key
                key = str(file_path.relative_to(self.docs_dir))
                self._cache[key] = doc

                logger.debug(f"Loaded: {doc.name} ({division}) - ~{approx_tokens} tokens")

            except Exception as e:
                logger.error(f"Failed to load {file_path}: {e}")

        self._loaded = True
        logger.info(f"Loaded {len(self._cache)} documents into cache")

    def get_docs_for_division(self, division: str) -> List[LoadedDoc]:
        """
        Get all documents for a division.

        Args:
            division: Division name (e.g., "warehouse", "hr")

        Returns:
            List of LoadedDoc for that division
        """
        self._load_all()

        division_lower = division.lower()
        return [
            doc for doc in self._cache.values()
            if doc.division == division_lower
        ]

    def get_all_docs(self) -> List[LoadedDoc]:
        """Get all loaded documents."""
        self._load_all()
        return list(self._cache.values())

    def get_stats(self) -> DocStats:
        """Get statistics about loaded documents."""
        self._load_all()

        by_division: Dict[str, Dict[str, int]] = {}
        total_chars = 0
        total_tokens = 0
        doc_list = []

        for doc in self._cache.values():
            if doc.division not in by_division:
                by_division[doc.division] = {
                    "docs": 0,
                    "chars": 0,
                    "tokens": 0,
                }

            by_division[doc.division]["docs"] += 1
            by_division[doc.division]["chars"] += doc.char_count
            by_division[doc.division]["tokens"] += doc.approx_tokens

            total_chars += doc.char_count
            total_tokens += doc.approx_tokens
            doc_list.append(f"{doc.division}/{doc.name}")

        return DocStats(
            total_docs=len(self._cache),
            total_chars=total_chars,
            total_tokens=total_tokens,
            by_division=by_division,
            doc_list=doc_list,
        )


# =============================================================================
# CONTEXT BUILDER
# =============================================================================

class DivisionContextBuilder:
    """
    Builds context strings for stuffing into LLM prompts.

    Division-aware: only includes docs relevant to user's division.
    Respects token limits by truncating or selecting subset.
    """

    def __init__(self, docs_dir_or_loader):
        """
        Initialize context builder.

        Args:
            docs_dir_or_loader: Either a Path to docs dir or a DocLoader instance
        """
        if isinstance(docs_dir_or_loader, DocLoader):
            self.loader = docs_dir_or_loader
        else:
            self.loader = DocLoader(Path(docs_dir_or_loader))

    def get_context_for_division(
        self,
        division: str,
        max_tokens: int = 200000,
        include_shared: bool = True,
    ) -> str:
        """
        Get formatted context string for a division.

        Args:
            division: User's division (e.g., "warehouse")
            max_tokens: Maximum tokens to include
            include_shared: Whether to include shared documents

        Returns:
            Formatted context string for prompt injection
        """
        # Get division docs
        docs = self.loader.get_docs_for_division(division)

        # Add shared docs if requested
        if include_shared and division != "shared":
            shared_docs = self.loader.get_docs_for_division("shared")
            docs.extend(shared_docs)

        if not docs:
            logger.warning(f"No documents found for division: {division}")
            return ""

        # Sort by token count (smaller docs first for better fit)
        docs.sort(key=lambda d: d.approx_tokens)

        # Build context respecting token limit
        sections = []
        tokens_used = 0
        docs_included = 0

        # Header
        header = f"=== COMPANY DOCUMENTATION ({division.upper()}) ===\n"
        header += "The following documents are your authoritative source for procedures and policies.\n"
        header += "Cite document names when answering questions.\n\n"

        header_tokens = len(header) // 4
        tokens_used += header_tokens
        sections.append(header)

        for doc in docs:
            # Check if we have room for this doc
            if tokens_used + doc.approx_tokens > max_tokens:
                # Try to fit a truncated version
                remaining_tokens = max_tokens - tokens_used
                if remaining_tokens > 500:  # Worth including a truncated version
                    max_chars = remaining_tokens * 4
                    truncated = doc.content[:max_chars]
                    truncated += "\n\n[DOCUMENT TRUNCATED - ASK FOR SPECIFIC SECTIONS]"

                    section = f"--- {doc.name} (from {doc.division}) ---\n\n"
                    section += truncated
                    section += "\n\n"

                    sections.append(section)
                    docs_included += 1

                break

            # Add full document
            section = f"--- {doc.name} (from {doc.division}) ---\n\n"
            section += doc.content
            section += "\n\n"

            sections.append(section)
            tokens_used += doc.approx_tokens
            docs_included += 1

        # Footer
        footer = f"=== END DOCUMENTATION ({docs_included} documents, ~{tokens_used} tokens) ===\n"
        sections.append(footer)

        context = "".join(sections)
        logger.info(f"Built context for {division}: {docs_included} docs, ~{tokens_used} tokens")

        return context

    def get_context_for_divisions(
        self,
        divisions: List[str],
        max_tokens: int = 200000,
    ) -> str:
        """
        Get context for multiple divisions.

        Useful for managers who need access to multiple areas.
        """
        all_docs = []

        for division in divisions:
            docs = self.loader.get_docs_for_division(division)
            all_docs.extend(docs)

        # Deduplicate (in case "shared" was included multiple times)
        seen_paths = set()
        unique_docs = []
        for doc in all_docs:
            if str(doc.path) not in seen_paths:
                seen_paths.add(str(doc.path))
                unique_docs.append(doc)

        # Sort and build
        unique_docs.sort(key=lambda d: d.approx_tokens)

        sections = []
        tokens_used = 0
        docs_included = 0

        header = f"=== COMPANY DOCUMENTATION (MULTI-DIVISION ACCESS) ===\n"
        header += f"Divisions: {', '.join(divisions)}\n"
        header += "Cite document names when answering questions.\n\n"

        header_tokens = len(header) // 4
        tokens_used += header_tokens
        sections.append(header)

        for doc in unique_docs:
            if tokens_used + doc.approx_tokens > max_tokens:
                break

            section = f"--- {doc.name} [{doc.division}] ---\n\n"
            section += doc.content
            section += "\n\n"

            sections.append(section)
            tokens_used += doc.approx_tokens
            docs_included += 1

        footer = f"=== END DOCUMENTATION ({docs_included} documents) ===\n"
        sections.append(footer)

        return "".join(sections)


# =============================================================================
# CLI
# =============================================================================

if __name__ == "__main__":
    import sys

    print("DocLoader v2.0.0 - Multi-format support")
    print(f"Supported extensions: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")
    print(f"DOCX support: {'Yes' if DOCX_AVAILABLE else 'No (install python-docx)'}")
    print(f"Excel/CSV support: {'Yes' if PANDAS_AVAILABLE else 'Limited (install pandas for better formatting)'}")
    print()

    # Default to ./manuals
    docs_dir = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("./manuals")

    if not docs_dir.exists():
        print(f"Directory not found: {docs_dir}")
        sys.exit(1)

    print(f"Loading documents from: {docs_dir}")
    print("=" * 60)

    loader = DocLoader(docs_dir)
    stats = loader.get_stats()

    print(f"\nTotal documents: {stats.total_docs}")
    print(f"Total characters: {stats.total_chars:,}")
    print(f"Approximate tokens: {stats.total_tokens:,}")

    print("\nBy division:")
    for division, div_stats in stats.by_division.items():
        print(f"  {division}:")
        print(f"    Docs: {div_stats['docs']}")
        print(f"    Tokens: {div_stats['tokens']:,}")

    print("\nDocument list:")
    for doc_name in stats.doc_list:
        print(f"  - {doc_name}")

    # Test context building
    if "--context" in sys.argv:
        builder = DivisionContextBuilder(loader)

        # Try warehouse division
        if "warehouse" in stats.by_division:
            print("\n" + "=" * 60)
            print("SAMPLE CONTEXT (warehouse, 50K tokens max)")
            print("=" * 60)
            context = builder.get_context_for_division("warehouse", max_tokens=50000)
            print(context[:2000] + "\n...[truncated]...")
