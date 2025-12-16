"""
Bulk Manual Uploader - Reads .docx and .json files and inserts into Azure PostgreSQL

Migrated from Supabase to direct Azure PostgreSQL connection via psycopg2.

Usage:
    python upload_manuals.py                    # Dry run
    python upload_manuals.py --execute          # Actually upload
    python upload_manuals.py --init-db          # Create tables first
    python upload_manuals.py --init-db --execute # Create tables + upload

Requirements:
    pip install python-docx psycopg2-binary python-dotenv

Folder structure expected:
    Manuals/Driscoll/
    +-- Warehouse/
    |   +-- Driver Check-in Manual.docx
    |   +-- Driver Check-in Manual.json  <- JSON also supported
    |   +-- Receiving Procedures.docx
    |   +-- ...
    +-- Sales/
    |   +-- ...
    +-- Credit/
    |   +-- ...
    +-- ...

Supported formats:
    .docx - Converted to markdown (headings, lists, tables preserved)
    .json - Pretty-printed JSON string (LLM reads it fine)
"""

import os
import json
import uuid
from pathlib import Path
from datetime import datetime
from docx import Document
import psycopg2
from psycopg2.extras import RealDictCursor
from dotenv import load_dotenv

load_dotenv()

# =============================================================================
# DATABASE CONFIG - Direct Azure PostgreSQL
# =============================================================================

DB_CONFIG = {
    "user": os.getenv("AZURE_PG_USER", "Mhartigan"),
    "password": os.getenv("AZURE_PG_PASSWORD", "Lalamoney3!"),
    "host": os.getenv("AZURE_PG_HOST", "enterprisebot.postgres.database.azure.com"),
    "port": int(os.getenv("AZURE_PG_PORT", "5432")),
    "database": os.getenv("AZURE_PG_DATABASE", "postgres"),
    "sslmode": "require"
}

# Schema for enterprise tables
SCHEMA = "enterprise"

# =============================================================================
# LOCAL CONFIG
# =============================================================================

# Path to your manuals folder (update for your environment)
MANUALS_ROOT = Path(os.getenv(
    "MANUALS_PATH",
    r"C:\Users\mthar\projects\enterprise_bot\Manuals\Driscoll"
))

# Department folder name -> slug mapping
# These will be created in the departments table if they don't exist
DEPARTMENT_SLUGS = {
    "warehouse": "warehouse",
    "warehouse operations": "warehouse",
    "sales": "sales",
    "credit": "credit",
    "credit department": "credit",
    "purchasing": "purchasing",
    "transportation": "transportation",
    "executive": "executive",
    "executive team": "executive",
    "hr": "executive",  # Map HR to executive if no HR dept
}

# =============================================================================
# DATABASE CONNECTION
# =============================================================================

def get_connection():
    """Create a new database connection."""
    return psycopg2.connect(**DB_CONFIG)


def init_database():
    """
    Create the enterprise schema and tables if they don't exist.
    Safe to run multiple times - uses IF NOT EXISTS.
    """
    conn = get_connection()
    cur = conn.cursor()
    
    try:
        # Create schema
        cur.execute(f"CREATE SCHEMA IF NOT EXISTS {SCHEMA};")
        
        # Create departments table
        cur.execute(f"""
            CREATE TABLE IF NOT EXISTS {SCHEMA}.departments (
                id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                name VARCHAR(100) NOT NULL,
                slug VARCHAR(50) UNIQUE NOT NULL,
                description TEXT,
                active BOOLEAN DEFAULT TRUE,
                created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
            );
        """)
        
        # Create department_content table
        cur.execute(f"""
            CREATE TABLE IF NOT EXISTS {SCHEMA}.department_content (
                id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                department_id UUID NOT NULL REFERENCES {SCHEMA}.departments(id),
                content_type VARCHAR(50) DEFAULT 'manual',
                title VARCHAR(255) NOT NULL,
                content TEXT NOT NULL,
                version INTEGER DEFAULT 1,
                active BOOLEAN DEFAULT TRUE,
                created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(department_id, title)
            );
        """)
        
        # Create indexes for fast lookups
        cur.execute(f"""
            CREATE INDEX IF NOT EXISTS idx_dept_content_dept_id 
            ON {SCHEMA}.department_content(department_id);
        """)
        cur.execute(f"""
            CREATE INDEX IF NOT EXISTS idx_dept_content_active 
            ON {SCHEMA}.department_content(active) WHERE active = TRUE;
        """)
        
        # Seed default departments
        default_departments = [
            ("Warehouse", "warehouse", "Warehouse operations and logistics"),
            ("Sales", "sales", "Sales department procedures"),
            ("Credit", "credit", "Credit and accounts receivable"),
            ("Purchasing", "purchasing", "Purchasing and procurement"),
            ("Transportation", "transportation", "Transportation and delivery"),
            ("Executive", "executive", "Executive team and HR policies"),
        ]
        
        for name, slug, description in default_departments:
            cur.execute(f"""
                INSERT INTO {SCHEMA}.departments (name, slug, description)
                VALUES (%s, %s, %s)
                ON CONFLICT (slug) DO NOTHING;
            """, (name, slug, description))
        
        conn.commit()
        print("[OK] Database initialized successfully")
        print(f"     Schema: {SCHEMA}")
        print(f"     Tables: departments, department_content")
        
        # Show department IDs for reference
        cur.execute(f"SELECT id, name, slug FROM {SCHEMA}.departments ORDER BY name;")
        print("\n     Department IDs:")
        for row in cur.fetchall():
            print(f"       {row[2]}: {row[0]}")
        
    except Exception as e:
        conn.rollback()
        print(f"[ERROR] Database initialization failed: {e}")
        raise
    finally:
        cur.close()
        conn.close()


# =============================================================================
# DOCX READER
# =============================================================================

def read_docx(file_path: Path) -> str:
    """
    Extract text from a .docx file.
    Preserves paragraph structure with double newlines.
    """
    doc = Document(file_path)
    
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    
    return "\n\n".join(paragraphs)


def read_json_manual(file_path: Path) -> str:
    """
    Read JSON manual and convert to string for context stuffing.
    LLM reads pretty-printed JSON just fine.
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return json.dumps(data, indent=2)


def read_docx_as_markdown(file_path: Path) -> str:
    """
    Extract text from .docx and convert to basic markdown.
    Handles headings, bold, lists somewhat.
    """
    doc = Document(file_path)
    
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Check paragraph style for headings
        style_name = para.style.name.lower() if para.style else ""
        
        if "heading 1" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name:
            lines.append(f"### {text}")
        elif "list" in style_name or "bullet" in style_name:
            lines.append(f"- {text}")
        else:
            lines.append(text)
    
    # Tables as markdown
    for table in doc.tables:
        table_lines = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            table_lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                # Header separator
                table_lines.append("|" + "|".join(["---"] * len(cells)) + "|")
        lines.extend(table_lines)
        lines.append("")  # blank line after table
    
    return "\n\n".join(lines)


# =============================================================================
# POSTGRESQL UPLOADER
# =============================================================================

def get_department_id(cur, slug: str) -> str:
    """Get department UUID by slug, or None if not found."""
    cur.execute(
        f"SELECT id FROM {SCHEMA}.departments WHERE slug = %s AND active = TRUE",
        (slug,)
    )
    row = cur.fetchone()
    return str(row[0]) if row else None


def upload_manuals(dry_run: bool = True):
    """
    Walk through manual folders and upload to PostgreSQL.
    
    Args:
        dry_run: If True, just print what would be uploaded without actually inserting
    """
    if not MANUALS_ROOT.exists():
        print(f"[ERROR] Manuals folder not found: {MANUALS_ROOT}")
        print("        Set MANUALS_PATH environment variable or update MANUALS_ROOT in script")
        return
    
    conn = get_connection()
    cur = conn.cursor()
    
    uploaded = 0
    updated = 0
    skipped = 0
    errors = []
    
    print(f"Scanning: {MANUALS_ROOT}\n")
    
    try:
        for dept_folder in MANUALS_ROOT.iterdir():
            if not dept_folder.is_dir():
                continue
            
            # Normalize folder name to match our mapping
            folder_name = dept_folder.name.lower().strip()
            dept_slug = DEPARTMENT_SLUGS.get(folder_name)
            
            if not dept_slug:
                print(f"[SKIP] Unknown department folder: {dept_folder.name}")
                skipped += 1
                continue
            
            dept_id = get_department_id(cur, dept_slug)
            if not dept_id:
                print(f"[SKIP] No department_id for: {dept_slug}")
                skipped += 1
                continue
            
            print(f"[DEPT] {dept_folder.name} -> {dept_slug} ({dept_id[:8]}...)")

            # Process all .docx and .json files in this folder
            for file in dept_folder.glob("*"):
                # Skip temp files and unsupported formats
                if file.name.startswith("~$"):
                    continue

                # Determine file type and reader
                if file.suffix == ".docx":
                    reader = read_docx_as_markdown
                elif file.suffix == ".json":
                    reader = read_json_manual
                else:
                    continue  # Skip unsupported file types

                title = file.stem  # Filename without extension

                try:
                    content = reader(file)
                    
                    if not content.strip():
                        print(f"       [SKIP] Empty file: {file.name}")
                        skipped += 1
                        continue
                    
                    print(f"       [FILE] {title} ({len(content):,} chars)")
                    
                    if not dry_run:
                        # Check if this manual already exists
                        cur.execute(f"""
                            SELECT id, version FROM {SCHEMA}.department_content 
                            WHERE department_id = %s AND title = %s
                        """, (dept_id, title))
                        existing = cur.fetchone()
                        
                        if existing:
                            # UPDATE existing - bump version
                            record_id = existing[0]
                            old_version = existing[1]
                            new_version = old_version + 1
                            
                            cur.execute(f"""
                                UPDATE {SCHEMA}.department_content 
                                SET content = %s, 
                                    version = %s, 
                                    active = TRUE,
                                    updated_at = CURRENT_TIMESTAMP
                                WHERE id = %s
                            """, (content, new_version, record_id))
                            
                            print(f"              [UPDATE] v{old_version} -> v{new_version}")
                            updated += 1
                        else:
                            # INSERT new
                            cur.execute(f"""
                                INSERT INTO {SCHEMA}.department_content 
                                (department_id, content_type, title, content, version, active)
                                VALUES (%s, %s, %s, %s, %s, %s)
                            """, (dept_id, "manual", title, content, 1, True))
                            
                            print(f"              [INSERT] v1")
                            uploaded += 1
                    else:
                        uploaded += 1

                except Exception as e:
                    print(f"       [ERROR] {file.name}: {e}")
                    errors.append(f"{file.name}: {e}")
        
        if not dry_run:
            conn.commit()
        
    except Exception as e:
        conn.rollback()
        print(f"\n[ERROR] Upload failed: {e}")
        raise
    finally:
        cur.close()
        conn.close()
    
    # Summary
    print("\n" + "="*60)
    print(f"{'DRY RUN ' if dry_run else ''}SUMMARY:")
    print(f"  [+] {'Would insert' if dry_run else 'Inserted'}: {uploaded}")
    print(f"  [~] {'Would update' if dry_run else 'Updated'}: {updated}")
    print(f"  [-] Skipped: {skipped}")
    print(f"  [!] Errors: {len(errors)}")
    
    if errors:
        print("\nErrors:")
        for e in errors:
            print(f"  - {e}")
    
    if dry_run:
        print("\n[INFO] This was a DRY RUN. To actually upload, run:")
        print("       python upload_manuals.py --execute")


def list_content():
    """List all content currently in the database."""
    conn = get_connection()
    cur = conn.cursor(cursor_factory=RealDictCursor)
    
    try:
        cur.execute(f"""
            SELECT 
                d.name as department,
                dc.title,
                dc.version,
                dc.active,
                LENGTH(dc.content) as content_length,
                dc.updated_at
            FROM {SCHEMA}.department_content dc
            JOIN {SCHEMA}.departments d ON dc.department_id = d.id
            ORDER BY d.name, dc.title
        """)
        
        rows = cur.fetchall()
        
        if not rows:
            print("[INFO] No content in database yet")
            return
        
        print(f"\nContent in {SCHEMA}.department_content:")
        print("="*80)
        
        current_dept = None
        for row in rows:
            if row['department'] != current_dept:
                current_dept = row['department']
                print(f"\n[{current_dept}]")
            
            status = "[OK]" if row['active'] else "[INACTIVE]"
            print(f"  {status} {row['title']} (v{row['version']}, {row['content_length']:,} chars)")
        
        print(f"\nTotal: {len(rows)} documents")
        
    finally:
        cur.close()
        conn.close()


# =============================================================================
# MAIN
# =============================================================================

if __name__ == "__main__":
    import sys
    
    args = sys.argv[1:]
    
    if "--help" in args or "-h" in args:
        print(__doc__)
        print("\nCommands:")
        print("  --init-db     Create schema and tables (safe to run multiple times)")
        print("  --execute     Actually perform the upload (default is dry run)")
        print("  --list        List current content in database")
        sys.exit(0)
    
    if "--list" in args:
        list_content()
        sys.exit(0)
    
    if "--init-db" in args:
        print("[INIT] Initializing database...\n")
        init_database()
        print()
        
        # If only --init-db, exit here
        if "--execute" not in args:
            sys.exit(0)
    
    # Check for --execute flag
    if "--execute" in args:
        print("[EXEC] EXECUTING UPLOAD (not a dry run)\n")
        upload_manuals(dry_run=False)
    else:
        print("[DRY] DRY RUN MODE (use --execute to actually upload)\n")
        upload_manuals(dry_run=True)